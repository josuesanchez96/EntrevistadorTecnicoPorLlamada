"""
rag_engine.py
Motor de RAG: carga de documentos, chunking, embeddings y almacén vectorial FAISS.
"""

from __future__ import annotations

import io
import logging
from typing import Optional

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings

logger = logging.getLogger(__name__)

# ── Parámetros de chunking ────────────────────────────────────────────────────
CHUNK_SIZE = 600
CHUNK_OVERLAP = 80

# Prefijos para distinguir la fuente de cada fragmento
CV_PREFIX = "[CV]"
JD_PREFIX = "[VACANTE]"


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extrae texto de DOCX incluyendo párrafos y tablas."""
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)
    return "\n".join(parts)


def _extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Extrae texto plano de PDF, DOCX o TXT dado como bytes."""
    fname = filename.lower()

    if fname.endswith(".doc") and not fname.endswith(".docx"):
        raise ValueError(
            "El formato .doc (Word antiguo) no está soportado. "
            "Abre el archivo en Word y guárdalo como .docx "
            "(Archivo → Guardar como → Documento de Word)."
        )

    if fname.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        except Exception as exc:
            logger.warning("Error extrayendo PDF '%s': %s", filename, exc)
            return ""

    if fname.endswith(".docx"):
        try:
            return _extract_docx_text(file_bytes)
        except Exception as exc:
            logger.warning("Error extrayendo DOCX '%s': %s", filename, exc)
            return ""

    # Fallback: tratar como texto plano UTF-8
    try:
        return file_bytes.decode("utf-8", errors="replace")
    except Exception:
        return ""


def build_vector_store(
    cv_bytes: bytes,
    cv_filename: str,
    jd_text: str,
) -> tuple[FAISS, str, str]:
    """
    Construye un almacén vectorial FAISS para una sesión.

    Args:
        cv_bytes:    Contenido en bytes del archivo de CV subido.
        cv_filename: Nombre del archivo (para detectar tipo).
        jd_text:     Texto de la descripción de puesto (Job Description).

    Returns:
        Tupla (vector_store, cv_text, jd_text).
    """
    embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", " ", ""],
    )

    # ── Procesar CV ───────────────────────────────────────────────────────────
    cv_text = _extract_text_from_bytes(cv_bytes, cv_filename)
    if not cv_text.strip():
        raise ValueError(
            "No se pudo extraer texto del CV. Si es Word, usa formato .docx "
            "(no .doc) y asegúrate de que el documento tenga texto, no solo imágenes."
        )

    logger.info("CV extraído: %d caracteres desde '%s'.", len(cv_text), cv_filename)

    cv_chunks = splitter.split_text(cv_text)
    cv_docs_text = [f"{CV_PREFIX} {chunk}" for chunk in cv_chunks]
    logger.info("CV dividido en %d fragmentos.", len(cv_docs_text))

    # ── Procesar Vacante ──────────────────────────────────────────────────────
    jd_chunks = splitter.split_text(jd_text)
    jd_docs_text = [f"{JD_PREFIX} {chunk}" for chunk in jd_chunks]
    logger.info("Vacante dividida en %d fragmentos.", len(jd_docs_text))

    all_texts = cv_docs_text + jd_docs_text
    if not all_texts:
        raise ValueError("Los documentos están vacíos después del chunking.")

    # ── Crear FAISS ───────────────────────────────────────────────────────────
    vector_store = FAISS.from_texts(all_texts, embedding=embeddings)
    logger.info(
        "Almacén vectorial creado con %d documentos en total.", len(all_texts)
    )

    return vector_store, cv_text, jd_text


def similarity_search(
    vector_store: FAISS,
    query: str,
    k: int = 5,
    filter_prefix: Optional[str] = None,
) -> list[str]:
    """
    Realiza una búsqueda semántica en el almacén vectorial.

    Args:
        vector_store:   El índice FAISS de la sesión.
        query:          Consulta en lenguaje natural.
        k:              Número de resultados a retornar.
        filter_prefix:  Si se especifica ('[CV]' o '[VACANTE]'), filtra por fuente.

    Returns:
        Lista de textos de los fragmentos más relevantes.
    """
    results = vector_store.similarity_search(query, k=k * 2 if filter_prefix else k)

    texts = [doc.page_content for doc in results]

    if filter_prefix:
        texts = [t for t in texts if t.startswith(filter_prefix)]

    return texts[:k]
